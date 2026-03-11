import keyboard
import pyautogui
import pyperclip
from docx import Document
from docx.shared import Inches
import time
import os
from datetime import datetime
import sys
import threading
from pathlib import Path

home_dir = os.path.expanduser("~")
default_docxfile = os.path.join(home_dir, "Downloads", "Screenshots_project", "Snap-Save-URL-Screenshot-Automation(2-file-option)", "shots_public.docx")
default_docxfile2 = os.path.join(home_dir, "Downloads", "Screenshots_project", "Snap-Save-URL-Screenshot-Automation(2-file-option)", "shots_restricted.docx")
shotfile = os.path.join(home_dir, "Downloads", "Screenshots_project", "Snap-Save-URL-Screenshot-Automation(2-file-option)", "shots\shot.png")

# Default file paths
#default_docxfile = r"C:\Users\lefte\Downloads\Screenshots_project\Snap-Save-URL-Screenshot-Automation(2-file-option)\shots_public.docx"
#default_docxfile2 = r"C:\Users\lefte\Downloads\Screenshots_project\Snap-Save-URL-Screenshot-Automation(2-file-option)\shots_restricted.docx"
#shotfile = r"C:\Users\lefte\Downloads\Screenshots_project\Snap-Save-URL-Screenshot-Automation(2-file-option)\shots\shot.png"

# Ask user for optional custom file names
custom_public = input(f"Please type a name for File 1 or press 'Enter' to use the default name (default: {os.path.basename(default_docxfile)}): ").strip()
if custom_public and not custom_public.lower().endswith('.docx'):
    custom_public += '.docx'

custom_restricted = input(f"Please type a name for File 2 or press 'Enter' to use the default name (default: {os.path.basename(default_docxfile2)}): ").strip()
if custom_restricted and not custom_restricted.lower().endswith('.docx'):
    custom_restricted += '.docx'

docxfile = os.path.join(os.path.dirname(default_docxfile), custom_public) if custom_public else default_docxfile
docxfile2 = os.path.join(os.path.dirname(default_docxfile2), custom_restricted) if custom_restricted else default_docxfile2

# Ensure files exist and load as Document objects
if not os.path.exists(docxfile):
    doc_public = Document()
    doc_public.save(docxfile)
else:
    doc_public = Document(docxfile)

if not os.path.exists(docxfile2):
    doc_restricted = Document()
    doc_restricted.save(docxfile2)
else:
    doc_restricted = Document(docxfile2)

exit_flag = threading.Event()  # Flag to signal exit


def do_cap1():
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    shot = pyautogui.screenshot()
    shot.save(shotfile)
    doc_public.add_paragraph(f"Screenshot taken at {timestamp}:")
    doc_public.add_picture(shotfile, width=Inches(7))   # screenshot taken

    keyboard.press_and_release('ctrl+l')
    time.sleep(1)
    pyperclip.copy("")  # Clear clipboard before grabbing URL
    keyboard.press_and_release('ctrl+c')
    time.sleep(1)
    url = pyperclip.paste()
    print(f"--------------Captured URL: {url}")
    doc_public.add_paragraph(f"Captured URL: {url}")
    doc_public.save(docxfile)
    print(f'Screenshot and URL saved and appended to {os.path.basename(docxfile)}.')


def do_cap2():
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    shot = pyautogui.screenshot()
    shot.save(shotfile)
    doc_restricted.add_paragraph(f"Screenshot taken at {timestamp}:")
    doc_restricted.add_picture(shotfile, width=Inches(7))   # screenshot taken

    keyboard.press_and_release('ctrl+l')
    time.sleep(1)
    pyperclip.copy("")  # Clear clipboard before grabbing URL
    keyboard.press_and_release('ctrl+c')
    time.sleep(1)
    url = pyperclip.paste()
    print(f"--------------Captured URL: {url}")
    doc_restricted.add_paragraph(f"Captured URL: {url}")
    doc_restricted.save(docxfile2)
    print(f'Screenshot and URL saved and appended to {os.path.basename(docxfile2)}.')


def input_listener():
    while not exit_flag.is_set():
        user_input = input()
        if user_input.strip().lower() == "exit":
            print("Exit command received. Exiting script.")
            exit_flag.set()
            keyboard.unhook_all_hotkeys()


# Hotkeys
keyboard.add_hotkey('ctrl+q', do_cap1)
keyboard.add_hotkey('ctrl+y', do_cap2)

print("\nScript running. Press Ctrl+Q to take screenshot and grab URL (public).")
print("Script running. Press Ctrl+Y to take screenshot and grab URL (restricted).")
print("Type 'exit' (without quotes) and press Enter to exit safely.")

# Start input listener in a separate thread
threading.Thread(target=input_listener, daemon=True).start()

# Main loop to keep script alive
while not exit_flag.is_set():
    time.sleep(0.1)  # Short sleep to reduce CPU use

print("Cleanup done, script exiting.")
sys.exit(0)